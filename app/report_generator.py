"""
Assessment DOCX Report Generator — profissional.
Generates professional AWS Assessment document following the template structure:
  Escopo → Cenário Atual (por serviço) → Recomendações → Considerações Gerais →
  Pain Points → Oportunidades → Referências
"""
import io
import logging
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.assessment.schemas import Assessment, ServiceSection
from app.assessment.sections_catalog import SECTIONS_CATALOG
from app.config import get_settings

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# COLOR PALETTE
# ---------------------------------------------------------------------------
PINK       = RGBColor(0xFF, 0x67, 0x9A)   # section titles (pink headings)
NAVY       = RGBColor(0x0B, 0x19, 0x56)   # subsection titles / labels
GREEN      = RGBColor(0x18, 0x80, 0x38)   # resource IDs / positive values
DARK_GRAY  = RGBColor(0x33, 0x33, 0x33)   # body text
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)   # footers / secondary
RED        = RGBColor(0xDC, 0x26, 0x26)   # critical severity
ORANGE     = RGBColor(0xEA, 0x58, 0x0C)   # high severity
AMBER      = RGBColor(0xD9, 0x77, 0x06)   # medium / warning
BLUE       = RGBColor(0x1D, 0x4E, 0xD8)   # low severity / links


def _hex_to_rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ---------------------------------------------------------------------------
# FORMATTING HELPERS
# ---------------------------------------------------------------------------

def _font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _heading(doc, text, level=1, color=NAVY):
    """Pink for level-1 section titles, Navy for subsections."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    _font(run, size={1: 18, 2: 14, 3: 12}.get(level, 11), bold=True, color=color)
    return p


def _body(doc, text, size=11, color=DARK_GRAY, indent_cm=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    _font(run, size=size, color=color)
    return p


def _bullet(doc, text, size=10.5, color=DARK_GRAY, indent_cm=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    _font(run, size=size, color=color)
    return p


def _kv(doc, key, value, key_color=NAVY, val_color=DARK_GRAY):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{key}: ")
    _font(r1, bold=True, color=key_color)
    r2 = p.add_run(str(value) if value is not None else "N/A")
    _font(r2, color=val_color)
    return p


def _aws_id(doc, text):
    """Resource ID in monospace green, indented."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    _font(run, name="Courier New", size=9.5, color=GREEN)
    return p


def _layer_section(doc, layer_data: dict | None, level: int = 2):
    """
    Render a service-layer block from the narrative JSON.
    layer_data = {"titulo": "...", "bullets": ["...", ...]}
    """
    if not layer_data or not isinstance(layer_data, dict):
        return
    titulo  = layer_data.get("titulo", "")
    bullets = layer_data.get("bullets") or []
    if titulo:
        _heading(doc, titulo, level=level, color=NAVY)
    for b in bullets:
        _bullet(doc, str(b))


def _conta_section(doc, conta: dict):
    """Render a single Landing Zone account block."""
    nome     = conta.get("nome", "")
    funcao   = conta.get("funcao_principal", "")
    recursos = conta.get("recursos_comuns") or []
    beneficio = conta.get("beneficio", "")

    _heading(doc, nome, level=3, color=NAVY)
    if funcao:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run("Função Principal: ")
        _font(r1, bold=True, color=NAVY)
        r2 = p.add_run(funcao)
        _font(r2)

    if recursos:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("Recursos Comuns:")
        _font(r, bold=True, color=NAVY)
        for rec in recursos:
            _bullet(doc, str(rec), indent_cm=0.5)

    if beneficio:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r1 = p.add_run("Benefício: ")
        _font(r1, bold=True, color=NAVY)
        r2 = p.add_run(beneficio)
        _font(r2, italic=True)


# ---------------------------------------------------------------------------
# MAIN GENERATOR
# ---------------------------------------------------------------------------

def generate_assessment_docx(report, client, assessment: Assessment) -> io.BytesIO:
    """
    Generate a professional AWS Assessment DOCX.
    assessment: pre-built Assessment object from the orchestrator.
    Returns a BytesIO buffer ready to stream as a download.
    """
    _settings = get_settings()
    pink = _hex_to_rgb(_settings.primary_color_hex) if _settings.primary_color_hex else PINK
    navy = _hex_to_rgb(_settings.accent_color_hex) if _settings.accent_color_hex else NAVY
    brand_name = _settings.consultancy_name or _settings.product_name

    # ── unpack raw data for fallbacks ─────────────────────────────────────
    raw        = report.raw_data or {}
    iam        = report.iam_findings or {}
    s3_data    = report.s3_findings or {}
    network    = report.network_findings or {}
    encryption = report.encryption_findings or {}
    unused     = report.unused_resources or {}
    ri_recs    = report.ri_recommendations or {}

    analysis_date = (report.completed_at or report.started_at or datetime.utcnow()).strftime("%d/%m/%Y")
    regions       = client.aws_regions or "us-east-1"
    savings_total = report.potential_savings or 0.0

    s3_costs_raw = s3_data.get("costs", {})
    s3_sec_raw   = s3_data.get("security", {})
    s3_buckets   = s3_costs_raw.get("buckets", [])
    public_bkts  = s3_sec_raw.get("public_buckets", [])
    no_lifecycle = sum(1 for b in s3_buckets if not b.get("has_lifecycle"))
    no_mfa       = iam.get("users_without_mfa", [])
    old_keys     = iam.get("users_with_old_keys", [])
    open_sgs     = network.get("open_security_groups", [])
    vpcs_no_flow = network.get("vpcs_without_flow_logs", [])
    no_guardduty = network.get("regions_without_guardduty", [])
    enc_ebs      = encryption.get("unencrypted_ebs", [])
    unattached   = unused.get("unattached_volumes", [])
    sp_recs      = ri_recs.get("savings_plans", [])
    rightsizing  = raw.get("rightsizing", {}).get("ec2", [])

    # ── build the DOCX ────────────────────────────────────────────────────
    doc = Document()
    doc_section = doc.sections[0]
    doc_section.page_height   = Cm(29.7)
    doc_section.page_width    = Cm(21.0)
    doc_section.left_margin   = Cm(2.5)
    doc_section.right_margin  = Cm(2.5)
    doc_section.top_margin    = Cm(2.5)
    doc_section.bottom_margin = Cm(2.5)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # ═════════════════════════════════════════════════════════════════════
    # CAPA
    # ═════════════════════════════════════════════════════════════════════
    if _settings.logo_path and os.path.isfile(_settings.logo_path):
        try:
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.add_run().add_picture(_settings.logo_path, width=Cm(5))
            doc.add_paragraph()
        except Exception:
            pass  # invalid logo file — skip silently

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Assessment Report")
    _font(r, size=28, bold=True, color=pink)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Arquitetura Cloud — {client.company}")
    _font(r, size=18, bold=True, color=navy)

    doc.add_paragraph()
    doc.add_paragraph()

    for label, value in [
        ("Versão:", "01"),
        ("Responsável Técnico:", brand_name),
        ("Data:", analysis_date),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label} ")
        _font(r1, size=12, color=LIGHT_GRAY)
        r2 = p.add_run(value)
        _font(r2, size=12, bold=True, color=navy)

    if client.aws_account_id:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"Conta AWS: {client.aws_account_id}")
        _font(r, name="Courier New", size=11, color=GREEN)

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════
    # SUMÁRIO
    # ═════════════════════════════════════════════════════════════════════
    _heading(doc, "Sumário", level=1, color=pink)

    # toc_entries: (main_title, sub_title, indent_level)
    # indent_level: 0=bold section, 1=account heading, 2=service subsection
    toc_entries: list[tuple[str, str, int]] = [
        ("Escopo", "", 0),
        ("Cenário Atual", "", 0),
    ]

    if assessment.secoes_por_conta:
        # Multi-account TOC: Conta → Serviço
        for acct in assessment.secoes_por_conta:
            toc_entries.append(("", f"Conta {acct.account_label}", 1))
            for spec in SECTIONS_CATALOG:
                if spec["slug"] in acct.sections:
                    toc_entries.append(("", spec["title"], 2))
    else:
        # Single-account TOC: flat list of services
        for spec in SECTIONS_CATALOG:
            if spec["slug"] in assessment.secoes_servicos:
                toc_entries.append(("", spec["title"], 1))

    toc_entries.append(("Recomendações de Serviços", "", 0))
    toc_entries.append(("Considerações Gerais", "", 0))
    if assessment.pain_points:
        toc_entries.append(("Pain Points", "", 0))
    if assessment.oportunidades:
        toc_entries.append(("Oportunidades", "", 0))
    toc_entries.append(("Referências", "", 0))

    for main, sub, indent in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        if main:
            r = p.add_run(main)
            _font(r, bold=True, color=navy)
        elif indent == 1:
            p.paragraph_format.left_indent = Cm(1.0)
            r = p.add_run(sub)
            _font(r, bold=True, color=pink)
        else:
            p.paragraph_format.left_indent = Cm(2.0)
            r = p.add_run(sub)
            _font(r, color=DARK_GRAY)

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════
    # ESCOPO
    # ═════════════════════════════════════════════════════════════════════
    _heading(doc, "Escopo", level=1, color=pink)
    if assessment.escopo_intro:
        _body(doc, assessment.escopo_intro)
    else:
        _body(doc, (
            f"O presente relatório visa a avaliação e identificação de possíveis melhorias "
            f"de segurança e boas práticas no ambiente AWS utilizado por {client.company}. "
            f"Ao longo do documento, serão apresentadas as recomendações necessárias e sugestões "
            f"de novos serviços, com suas respectivas finalidades, que podem ser implementadas."
        ))
        doc.add_paragraph()
        _body(doc, (
            "Este trabalho é baseado no AWS Well-Architected Framework, publicamente reconhecido "
            "como o conjunto das melhores práticas para a implementação de aplicações em nuvem."
        ))

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════
    # CENÁRIO ATUAL
    # ═════════════════════════════════════════════════════════════════════
    _heading(doc, "Cenário Atual", level=1, color=pink)

    if assessment.cenario_intro:
        _body(doc, assessment.cenario_intro)
    else:
        acct = client.aws_account_id or raw.get("account_id", "N/A")
        _body(doc, (
            f"Realizamos o levantamento técnico da conta {acct}, onde {client.company} opera "
            f"sua infraestrutura AWS. Foram coletados dados de segurança, custo e conformidade "
            f"nas regiões: {regions}."
        ))

    if assessment.vulnerabilidades_criticas:
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run("Vulnerabilidades Críticas Identificadas:")
        _font(r, bold=True, color=RED)
        for vuln in assessment.vulnerabilidades_criticas:
            _bullet(doc, str(vuln), color=RED)

    if assessment.arquitetura_atual:
        _heading(doc, "Arquitetura", level=2, color=navy)
        _body(doc, assessment.arquitetura_atual)

    # ── service sections ──────────────────────────────────────────────────
    if assessment.secoes_por_conta:
        # Multi-account: group by account, service sections at level 3
        for acct in assessment.secoes_por_conta:
            _heading(doc, f"Conta {acct.account_label}", level=2, color=pink)
            for spec in SECTIONS_CATALOG:
                svc_section = acct.sections.get(spec["slug"])
                if not svc_section or not svc_section.presente:
                    continue
                _render_service_section(doc, spec["title"], svc_section,
                                        navy_color=navy, level=3)
    else:
        # Single-account: flat service sections at level 2
        for spec in SECTIONS_CATALOG:
            svc_section = assessment.secoes_servicos.get(spec["slug"])
            if not svc_section or not svc_section.presente:
                continue
            _render_service_section(doc, spec["title"], svc_section,
                                    navy_color=navy, level=2)

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════
    # RECOMENDAÇÕES DE SERVIÇOS
    # ═════════════════════════════════════════════════════════════════════
    _heading(doc, "Recomendações de Serviços", level=1, color=pink)
    _body(doc, (
        "Nesta seção, listamos alguns serviços que poderão ser utilizados futuramente "
        f"no ambiente de {client.company}. O objetivo é apresentar novas soluções que "
        "contribuam para a estabilidade e segurança do ambiente."
    ))

    _heading(doc, "Proposta de Estrutura das Contas e Landing Zone", level=2, color=navy)

    lz = assessment.landing_zone
    if lz is not None:
        if not lz.usar_organizations and lz.contas:
            for conta in lz.contas:
                _conta_section(doc, {
                    "nome": conta.nome,
                    "funcao_principal": conta.funcao_principal,
                    "recursos_comuns": conta.recursos_comuns,
                    "beneficio": conta.beneficio,
                })
        elif lz.usar_organizations and lz.otimizacao_governanca:
            _heading(doc, "Otimização da Governança Multi-Conta", level=3, color=navy)
            for item in lz.otimizacao_governanca:
                _bullet(doc, str(item))
        else:
            _render_generic_landing_zone(doc, navy)
    else:
        _render_generic_landing_zone(doc, navy)

    # Default AWS Backup recommendation
    _heading(doc, "AWS Backup", level=2, color=navy)
    _body(doc, (
        "O AWS Backup permite centralizar e automatizar políticas de backup para recursos "
        f"AWS no ambiente de {client.company}. A adoção correta do serviço permitiria:"
    ))
    _bullet(doc, "Padronizar políticas de backup para volumes EBS associados às instâncias EC2.")
    _bullet(doc, "Definir políticas de retenção conforme criticidade do workload.")
    _bullet(doc, "Simplificar a gestão operacional de backups com um único plano centralizado.")

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════
    # CONSIDERAÇÕES GERAIS
    # ═════════════════════════════════════════════════════════════════════
    _heading(doc, "Considerações Gerais", level=1, color=pink)

    consider = assessment.consideracoes
    if consider.introducao:
        _body(doc, consider.introducao)
    else:
        _body(doc, (
            f"A arquitetura atual de {client.company} na AWS possui pontos de melhoria em termos "
            f"de segurança, governança e eficiência de custos, principalmente por meio da "
            f"padronização e consolidação dos recursos existentes."
        ))

    if consider.organizacao_governanca:
        _heading(doc, "Organização e Governança", level=2, color=navy)
        for b in consider.organizacao_governanca:
            _bullet(doc, str(b))

    vpc_bullets = consider.vpc
    seg_bullets = consider.seguranca
    acc_bullets = consider.acesso

    if vpc_bullets or seg_bullets or acc_bullets:
        _heading(doc, "Rede, Segurança e Acesso", level=2, color=navy)

        if vpc_bullets:
            p = doc.add_paragraph()
            r = p.add_run("VPC")
            _font(r, bold=True, color=navy)
            for b in vpc_bullets:
                _bullet(doc, str(b), indent_cm=0.5)

        if seg_bullets:
            p = doc.add_paragraph()
            r = p.add_run("Segurança")
            _font(r, bold=True, color=navy)
            for b in seg_bullets:
                _bullet(doc, str(b), indent_cm=0.5)

        if acc_bullets:
            p = doc.add_paragraph()
            r = p.add_run("Acesso")
            _font(r, bold=True, color=navy)
            for b in acc_bullets:
                _bullet(doc, str(b), indent_cm=0.5)

    finops_bullets = consider.performance_finops
    if not finops_bullets:
        finops_bullets = []
        if savings_total > 0:
            finops_bullets.append(
                f"Economia potencial identificada de ${savings_total:.0f}/mês "
                f"({report.savings_percentage or 0:.1f}% de redução possível)."
            )
        if sp_recs:
            finops_bullets.append(
                "Savings Plans: aquisição de compromissos pode reduzir custos de computação em até 60%."
            )
        if no_lifecycle:
            finops_bullets.append(
                f"Implementar S3 Lifecycle Policies em {no_lifecycle} bucket(s) para reduzir custos de armazenamento."
            )
        if rightsizing:
            finops_bullets.append(
                f"Rightsizing de {len(rightsizing)} instância(s) EC2 superdimensionadas identificadas pelo Compute Optimizer."
            )

    if finops_bullets:
        _heading(doc, "Recomendações de Performance e FinOps", level=2, color=navy)
        for b in finops_bullets:
            _bullet(doc, str(b))

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════
    # PAIN POINTS
    # ═════════════════════════════════════════════════════════════════════
    _heading(doc, "Pain Points", level=1, color=pink)
    _body(doc, "Ações críticas e imediatas para execução após alinhamento:")

    if assessment.pain_points:
        for pp in assessment.pain_points:
            _bullet(doc, str(pp))
    else:
        if iam.get("root_mfa_enabled") is False:
            _bullet(doc, "Habilitar MFA na conta root imediatamente — risco crítico de segurança.")
        if no_mfa:
            _bullet(doc, f"Habilitar MFA para {len(no_mfa)} usuário(s) IAM sem autenticação multifator.")
        if old_keys:
            _bullet(doc, f"Rotacionar {len(old_keys)} chave(s) de acesso com mais de 90 dias.")
        if public_bkts:
            _bullet(doc, f"Bloquear acesso público em {len(public_bkts)} bucket(s) S3 expostos.")
        if open_sgs:
            _bullet(doc, f"Revisar {len(open_sgs)} Security Group(s) com portas abertas para 0.0.0.0/0.")
        if no_guardduty:
            _bullet(doc, f"Habilitar GuardDuty nas regiões sem monitoramento: {', '.join(no_guardduty[:5])}.")
        if vpcs_no_flow:
            _bullet(doc, f"Ativar VPC Flow Logs em {len(vpcs_no_flow)} VPC(s) para visibilidade de tráfego.")
        if unattached:
            _bullet(doc, f"Remover ou snapshottear {len(unattached)} volume(s) EBS desanexados gerando custo.")
        if no_lifecycle:
            _bullet(doc, f"Implementar S3 Lifecycle Policy em {no_lifecycle} bucket(s) sem política de ciclo de vida.")
        if enc_ebs:
            _bullet(doc, f"Habilitar criptografia em {len(enc_ebs)} volume(s) EBS sem criptografia em repouso.")

    # ═════════════════════════════════════════════════════════════════════
    # OPORTUNIDADES
    # ═════════════════════════════════════════════════════════════════════
    if assessment.oportunidades or savings_total > 0 or sp_recs or unattached:
        doc.add_paragraph()
        _heading(doc, "Oportunidades", level=1, color=pink)
        _body(doc, "Oportunidades para um projeto de modernização e refactoring de infraestrutura:")

        if assessment.oportunidades:
            for opp in assessment.oportunidades:
                _bullet(doc, str(opp))
        else:
            if savings_total > 0:
                _bullet(doc, f"Implementar otimizações identificadas para reduzir ${savings_total:.0f}/mês ({report.savings_percentage or 0:.1f}% do custo atual).")
            if sp_recs:
                _bullet(doc, "Expandir cobertura de Savings Plans para reduzir exposição ao preço On-Demand.")
            if unattached:
                _bullet(doc, f"Eliminar {len(unattached)} volume(s) EBS ocioso(s) e revisar política de snapshots.")
            if no_lifecycle:
                _bullet(doc, "Configurar políticas de S3 Lifecycle para logs (S3 Intelligent-Tiering e Glacier Instant Retrieval).")
            if rightsizing:
                _bullet(doc, f"Aplicar recomendações de rightsizing do Compute Optimizer em {len(rightsizing)} instância(s).")
            _bullet(doc, "Implementar governança multi-conta através de AWS Organizations e Landing Zone.")
            _bullet(doc, "Adotar AWS Backup para orquestração centralizada de snapshots EBS e RDS.")

        doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════
    # REFERÊNCIAS
    # ═════════════════════════════════════════════════════════════════════
    _heading(doc, "Referências", level=1, color=pink)
    refs = [
        ("AWS Well-Architected Framework",
         "https://docs.aws.amazon.com/wellarchitected/latest/framework/welcome.html"),
        ("IAM Best Practices",
         "https://docs.aws.amazon.com/IAM/latest/UserGuide/best-practices.html"),
        ("AWS Backup",
         "https://docs.aws.amazon.com/aws-backup/latest/devguide/whatisbackup.html"),
        ("Landing Zones",
         "https://docs.aws.amazon.com/controltower/latest/userguide/landing-zone.html"),
        ("AWS Organizations",
         "https://docs.aws.amazon.com/organizations/latest/userguide/orgs_introduction.html"),
        ("Amazon RDS",
         "https://docs.aws.amazon.com/AmazonRDS/latest/UserGuide/Welcome.html"),
        ("AWS Savings Plans",
         "https://docs.aws.amazon.com/savingsplans/latest/userguide/what-is-savings-plans.html"),
        ("Amazon GuardDuty",
         "https://docs.aws.amazon.com/guardduty/latest/ug/what-is-guardduty.html"),
        ("AWS Security Hub",
         "https://docs.aws.amazon.com/securityhub/latest/userguide/what-is-securityhub.html"),
        ("AWS Cost Explorer",
         "https://docs.aws.amazon.com/cost-management/latest/userguide/ce-what-is.html"),
    ]
    for name, url in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r1 = p.add_run(f"• {name}: ")
        _font(r1, bold=True, color=navy)
        r2 = p.add_run(url)
        _font(r2, color=BLUE)

    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_footer.add_run(
        f"Documento gerado automaticamente em "
        f"{datetime.utcnow().strftime('%d/%m/%Y às %H:%M')} UTC"
    )
    _font(r, size=9, italic=True, color=LIGHT_GRAY)

    # ── save ──────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# SERVICE SECTION RENDERER
# ---------------------------------------------------------------------------

def _render_service_section(doc, title: str, section: ServiceSection, navy_color, level: int = 2):
    """Render one ServiceSection of the Assessment."""
    _heading(doc, title, level=level, color=navy_color)

    if section.intro:
        _body(doc, section.intro)

    if section.seguranca:
        p = doc.add_paragraph()
        r = p.add_run("Segurança")
        _font(r, bold=True, color=navy_color)
        for b in section.seguranca:
            _bullet(doc, str(b), indent_cm=0.5)

    if section.arquitetura:
        p = doc.add_paragraph()
        r = p.add_run("Arquitetura")
        _font(r, bold=True, color=navy_color)
        for b in section.arquitetura:
            _bullet(doc, str(b), indent_cm=0.5)

    if section.guide_list:
        p = doc.add_paragraph()
        r = p.add_run("Guide List")
        _font(r, bold=True, color=navy_color)
        for item in section.guide_list:
            _bullet(doc, str(item.text), indent_cm=0.5)

    if section.melhorias:
        p = doc.add_paragraph()
        r = p.add_run("Melhorias")
        _font(r, bold=True, color=navy_color)
        for b in section.melhorias:
            _bullet(doc, str(b), indent_cm=0.5)


def _render_generic_landing_zone(doc, navy_color):
    """Fallback landing zone when assessment.landing_zone is None."""
    _body(doc, (
        "Recomenda-se a implementação de AWS Organizations com Landing Zones para "
        "segregar os ambientes em contas distintas: Management, Network, Development e Production."
    ))
    for nome, funcao, beneficio in [
        ("Conta Management (Gerenciamento)",
         "Centralizar e gerenciar todas as contas da AWS via AWS Organizations.",
         "Facilita o gerenciamento das contas AWS, introduzindo governança, organização e segurança."),
        ("Conta Network (Rede)",
         "Centralizar todos os recursos de rede e conectividade (Transit Gateway, VPN).",
         "Isola a complexidade da rede, aplicando políticas de segurança consistentes."),
        ("Conta Development (Desenvolvimento)",
         "Hospedar ambientes de desenvolvimento e testes não críticos.",
         "Permite inovação rápida sem risco ao ambiente de produção."),
        ("Conta Production (Produção)",
         "Hospedar aplicações e serviços críticos que atendem usuários finais.",
         "Garante o mais alto nível de segurança e isolamento para serviços de missão crítica."),
    ]:
        _heading(doc, nome, level=3, color=navy_color)
        _kv(doc, "Função Principal", funcao)
        _kv(doc, "Benefício", beneficio)
